"""
pdf_extractor.py
----------------
Extracts raw text from uploaded resumes (PDF or DOCX).

Uses:
    - PyMuPDF (fitz)  -> primary PDF extraction (fast, layout-aware)
    - pdfplumber      -> fallback PDF extraction (better for tricky layouts/tables)
    - python-docx     -> DOCX extraction
"""

import os
from typing import Optional


class ResumeTextExtractor:
    """Extracts plain text from a resume file (.pdf or .docx)."""

    SUPPORTED_EXTENSIONS = (".pdf", ".docx")

    def extract(self, file_path: str) -> str:
        """
        Extract raw text from a resume file.

        Args:
            file_path: path to a .pdf or .docx file.

        Returns:
            Extracted plain text as a single string.

        Raises:
            ValueError: if the file extension is unsupported.
            FileNotFoundError: if the file does not exist.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file type '{ext}'. Supported types: {self.SUPPORTED_EXTENSIONS}"
            )

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from a PDF, trying PyMuPDF first, then pdfplumber as a fallback."""
        text = self._extract_pdf_pymupdf(file_path)
        if text and text.strip():
            return text

        text = self._extract_pdf_pdfplumber(file_path)
        return text or ""

    @staticmethod
    def _extract_pdf_pymupdf(file_path: str) -> Optional[str]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return None

        try:
            doc = fitz.open(file_path)
            text_parts = [page.get_text() for page in doc]
            doc.close()
            return "\n".join(text_parts)
        except Exception:
            return None

    @staticmethod
    def _extract_pdf_pdfplumber(file_path: str) -> Optional[str]:
        try:
            import pdfplumber
        except ImportError:
            return None

        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception:
            return None

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "python-docx is required for DOCX extraction. Install with: "
                "pip install python-docx"
            ) from e

        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also pull text out of tables (some resumes use table layouts)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)


if __name__ == "__main__":
    # Simple manual smoke test
    extractor = ResumeTextExtractor()
    print("ResumeTextExtractor ready. Supported types:", extractor.SUPPORTED_EXTENSIONS)
